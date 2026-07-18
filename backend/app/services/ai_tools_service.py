"""
AI Tools Service — PDF text extraction, summarization, key points, title generation,
sentiment analysis, and report generation.

Uses:
- pymupdf (fitz) for high-quality PDF text extraction
- HuggingFace transformers for NLP (bart-large-cnn, sentiment-analysis)
- python-docx for report generation
"""

import io
import os
import re
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ── Lazy-loaded heavy imports ────────────────────────────────────────
import fitz  # pymupdf — always available (lightweight PDF parsing)
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cloud LLM (Groq) — primary AI engine when configured. All public functions
# return Optional and we fall back to local HuggingFace pipelines on None.
from app.services import cloud_ai_service

logger = logging.getLogger("ai_tools")

# HuggingFace pipelines will be lazily loaded per-call to avoid blocking
# the event loop on first import (model download / warm-up can take time).
# For production, pre-load in a startup event.

_SUMMARIZER = None
_SENTIMENT_PIPELINE = None

MODEL_SUMMARY = "facebook/bart-large-cnn"  # good summarization model
CHUNK_MAX_CHARS = 1024  # Safe token limit for bart-large-cnn
OCR_MAX_PAGES = 30  # Cap OCR work on large scanned PDFs to keep requests responsive


# ── Lazy initializers ────────────────────────────────────────────────

def _get_summarizer():
    """Lazily load the summarization pipeline."""
    global _SUMMARIZER
    if _SUMMARIZER is None:
        from transformers import pipeline
        logger.info("Loading summarization model: %s", MODEL_SUMMARY)
        _SUMMARIZER = pipeline("summarization", model=MODEL_SUMMARY)
    return _SUMMARIZER


def _get_sentiment_pipeline():
    """Lazily load the sentiment-analysis pipeline."""
    global _SENTIMENT_PIPELINE
    if _SENTIMENT_PIPELINE is None:
        from transformers import pipeline
        logger.info("Loading sentiment-analysis model")
        _SENTIMENT_PIPELINE = pipeline("sentiment-analysis")
    return _SENTIMENT_PIPELINE


# ── Text extraction ──────────────────────────────────────────────────

def extract_pdf_text(pdf_bytes: bytes) -> Tuple[str, int, bool]:
    """Extract full text from a PDF using pymupdf.

    Falls back to OCR (Tesseract via PyMuPDF) for scanned / image-only pages
    that carry no embedded text layer, so scanned documents still produce a
    usable analysis instead of an empty report.

    Returns:
        (extracted_text, page_count, ocr_used)
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_count = doc.page_count
    full_text_parts: List[str] = []

    for page_num in range(page_count):
        page = doc[page_num]
        page_text = page.get_text("text")  # plain text with layout
        full_text_parts.append(page_text)

    doc.close()
    full_text = "\n\n".join(full_text_parts)
    logger.info("Extracted %d chars of text from %d pages", len(full_text), page_count)

    # Scanned / image-only PDF: no usable embedded text — try OCR.
    if len(full_text.strip()) < 20:
        logger.info("Embedded text negligible; attempting OCR fallback")
        ocr_text = _extract_pdf_text_ocr(pdf_bytes)
        if ocr_text.strip():
            logger.info("OCR fallback extracted %d chars", len(ocr_text))
            return ocr_text, page_count, True

    return full_text, page_count, False


def _extract_pdf_text_ocr(pdf_bytes: bytes) -> str:
    """OCR every page of a scanned PDF and return the recognized text.

    Uses PyMuPDF's built-in Tesseract bridge with a Hindi + English model so
    both scripts are recognized. Returns an empty string if OCR is unavailable
    (e.g. Tesseract not installed) so callers can degrade gracefully.
    """
    # Point PyMuPDF at the Tesseract language data. Prefer an existing env var;
    # otherwise fall back to the standard Windows install location.
    tessdata = os.environ.get("TESSDATA_PREFIX")
    if not tessdata:
        default = r"C:\Program Files\Tesseract-OCR\tessdata"
        if os.path.isdir(default):
            tessdata = default
            os.environ["TESSDATA_PREFIX"] = tessdata
    if not tessdata or not os.path.isdir(tessdata):
        logger.warning("OCR unavailable: TESSDATA_PREFIX not found")
        return ""

    # Choose the richest language model actually present.
    have_hin = os.path.isfile(os.path.join(tessdata, "hin.traineddata"))
    have_eng = os.path.isfile(os.path.join(tessdata, "eng.traineddata"))
    if have_hin and have_eng:
        language = "hin+eng"
    elif have_hin:
        language = "hin"
    else:
        language = "eng"

    parts: List[str] = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            # Cap OCR pages: analysis only needs enough text to summarize, and
            # OCR is expensive (~1s+/page). This keeps large scans responsive.
            max_ocr_pages = min(doc.page_count, OCR_MAX_PAGES)
            for page_num in range(max_ocr_pages):
                page = doc[page_num]
                try:
                    tp = page.get_textpage_ocr(
                        language=language, dpi=200, full=True
                    )
                    parts.append(page.get_text("text", textpage=tp))
                except Exception as e:  # per-page OCR failure shouldn't abort all
                    logger.warning("OCR failed on a page: %s", e)
        finally:
            doc.close()
    except Exception as e:
        logger.warning("OCR fallback failed: %s", e)
        return ""

    return "\n\n".join(parts)


def _compute_reading_time(word_count: int) -> str:
    """Estimate reading time in minutes (average 200 wpm)."""
    minutes = max(1, round(word_count / 200))
    if minutes < 1:
        return "< 1 minute"
    elif minutes == 1:
        return "1 minute"
    else:
        return f"{minutes} minutes"


# ── OCR quality heuristic ─────────────────────────────────────────────

def _looks_like_prose(text: str) -> bool:
    """Best-effort check that text reads like real prose, not OCR garbage.

    Scanned OCR (especially Devanagari) frequently collapses word spacing,
    producing long space-free runs like ``मध्यप्रदेशशासननगरीयविकास``. Echoing
    that back as a "summary" or "title" looks broken, so callers use this to
    decide whether a raw-text fallback is safe. This only gates the local
    last-resort fallbacks — the cloud LLM handles messy OCR fine on its own.
    """
    sample = text.strip()
    if len(sample) < 40:
        return False

    tokens = sample.split()
    if not tokens:
        return False

    # Ratio of whitespace to total length. Normal prose sits well above 8%;
    # space-collapsed OCR drops far below that.
    space_ratio = sample.count(" ") / len(sample)
    if space_ratio < 0.06:
        return False

    # Average token length. Real words average ~4-7 chars; joined OCR runs
    # push this much higher.
    avg_token_len = sum(len(t) for t in tokens) / len(tokens)
    if avg_token_len > 14:
        return False

    return True


# ── Text chunking for models with token limits ───────────────────────

def _chunk_text(text: str, max_chars: int = CHUNK_MAX_CHARS) -> List[str]:
    """Split text into roughly equal chunks, respecting sentence boundaries."""
    if len(text) <= max_chars:
        return [text]

    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks: List[str] = []
    current_chunk: List[str] = []
    current_len = 0

    for sentence in sentences:
        if current_len + len(sentence) > max_chars and current_chunk:
            chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]
            current_len = len(sentence)
        else:
            current_chunk.append(sentence)
            current_len += len(sentence)

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks if chunks else [text]


# ── AI Analysis ──────────────────────────────────────────────────────

def _generate_summary(text: str, max_length: int = 150, min_length: int = 40) -> str:
    """Generate a concise summary using cloud LLM (Groq) with HuggingFace fallback."""
    # Try cloud LLM first — higher quality, faster
    cloud_result = cloud_ai_service.summarize(text, max_words=max_length // 2)
    if cloud_result:
        logger.info("Summary: cloud LLM (Groq)")
        return cloud_result

    # Fallback: HuggingFace BART (slower, lower quality)
    logger.info("Summary: HuggingFace fallback")
    # Truncate if text is too long
    input_text = text[:4000] if len(text) > 4000 else text
    if len(input_text.strip()) < 50:
        return input_text.strip()

    try:
        summarizer = _get_summarizer()
        result = summarizer(
            input_text,
            max_length=min(max_length, 200),
            min_length=min(min_length, 30),
            do_sample=False,
        )
        if result and len(result) > 0:
            return result[0]["summary_text"].strip()
    except Exception as e:
        logger.warning("Summarization failed: %s", e)

    # Last-resort fallback: first few sentences, but only if the text looks
    # like real prose. Scanned OCR often produces space-collapsed runs
    # (e.g. Devanagari words joined together) that read as garbage — surface
    # an honest message instead of dumping that back to the user.
    if not _looks_like_prose(input_text):
        return (
            "A clean summary could not be generated automatically — the "
            "document text appears to be low-quality or scanned output. "
            "See the key points below for the extracted highlights."
        )
    sentences = re.split(r'(?<=[.!?])\s+', input_text)
    return " ".join(sentences[:3])


def _generate_structured_summary(text: str) -> str:
    """Structured, labeled summary for the clean PDF report.

    Uses the cloud LLM's structured mode. Falls back to the plain prose summary
    (as a single 'Summary:' line) when the cloud engine is unavailable so the
    PDF renderer always has something sensible to lay out.
    """
    cloud_result = cloud_ai_service.summarize_structured(text)
    if cloud_result:
        logger.info("Structured summary: cloud LLM (Groq)")
        return cloud_result

    logger.info("Structured summary: prose fallback")
    return _generate_summary(text)


def _extract_key_points(text: str, num_points: int = 5) -> List[str]:
    """Extract key points using cloud LLM with HuggingFace + heuristic fallback."""
    if len(text.strip()) < 50:
        return ["Text too short for key point extraction."]

    # Try cloud LLM first — gives genuinely distinct insights, not duplicates
    cloud_points = cloud_ai_service.extract_key_points(text, num_points=num_points)
    if cloud_points:
        logger.info("KeyPoints: cloud LLM (Groq), %d points", len(cloud_points))
        return cloud_points

    # Fallback: HuggingFace per-chunk summarization
    logger.info("KeyPoints: HuggingFace fallback")
    chunks = _chunk_text(text, max_chars=1200)
    all_points: List[str] = []

    # Summarize each chunk to get a mini-insight
    for chunk in chunks[:5]:  # limit to avoid excessive API calls
        if len(chunk.strip()) < 20:
            continue
        try:
            mini_summary = _generate_summary(chunk, max_length=80, min_length=20)
            if mini_summary and mini_summary not in all_points:
                all_points.append(mini_summary)
        except Exception:
            pass

    if len(all_points) < num_points:
        # Last-resort: extract longest sentences as points
        sentences = re.split(r'(?<=[.!?])\s+', text)
        long_sentences = sorted(
            [s.strip() for s in sentences if len(s.strip()) > 30],
            key=len, reverse=True
        )
        for s in long_sentences:
            if s not in all_points and len(all_points) < num_points:
                all_points.append(s)

    return all_points[:num_points]


def _generate_title(text: str) -> str:
    """Generate a smart title using cloud LLM with HuggingFace fallback."""
    if len(text.strip()) < 30:
        return "Untitled Document"

    # Try cloud LLM first — produces clean, properly-formatted titles
    cloud_title = cloud_ai_service.generate_title(text)
    if cloud_title:
        logger.info("Title: cloud LLM (Groq)")
        return cloud_title

    # Fallback: HuggingFace BART
    logger.info("Title: HuggingFace fallback")
    try:
        summarizer = _get_summarizer()
        # Use a very short summary as title
        result = summarizer(
            text[:3000],
            max_length=30,
            min_length=5,
            do_sample=False,
        )
        if result and len(result) > 0:
            title = result[0]["summary_text"].strip()
            # Clean up: capitalize words, remove trailing periods
            title = title.rstrip(".")
            title = " ".join(
                w.capitalize() if len(w) > 2 else w for w in title.split()
            )
            return title
    except Exception as e:
        logger.warning("Title generation failed: %s", e)

    # Last-resort: use the first sentence, but never echo back
    # space-collapsed OCR garbage as a "title".
    if not _looks_like_prose(text):
        return "Document Analysis"
    sentences = re.split(r'(?<=[.!?])\s+', text)
    first = sentences[0].strip().rstrip(".")
    if len(first) > 60:
        first = first[:60] + "..."
    return first if first else "Document Analysis"


def _analyze_sentiment(text: str) -> Tuple[str, float]:
    """Analyze sentiment using cloud LLM with HuggingFace fallback."""
    if len(text.strip()) < 20:
        return "neutral", 50.0

    # Try cloud LLM first — better nuance than binary distilbert
    cloud_result = cloud_ai_service.analyze_sentiment(text)
    if cloud_result:
        logger.info("Sentiment: cloud LLM (Groq)")
        return cloud_result

    # Fallback: HuggingFace distilbert (positive/negative only, no neutral)
    logger.info("Sentiment: HuggingFace fallback")
    sample = text[:1000] if len(text) > 1000 else text
    try:
        pipeline = _get_sentiment_pipeline()
        result = pipeline(sample)
        if result and len(result) > 0:
            label = result[0]["label"].lower()
            score = round(result[0]["score"] * 100, 1)
            if label in ("positive", "pos", "5 stars", "4 stars"):
                return "positive", score
            elif label in ("negative", "neg", "1 star", "2 stars"):
                return "negative", score
            else:
                return "neutral", score
    except Exception as e:
        logger.warning("Sentiment analysis failed: %s", e)

    return "neutral", 50.0


# ── Word count ───────────────────────────────────────────────────────

def _count_words(text: str) -> int:
    """Count words in text."""
    return len(re.findall(r'\b\w+\b', text))


# ── Full analysis pipeline ───────────────────────────────────────────

def analyze_pdf(pdf_bytes: bytes) -> Dict:
    """Run the complete AI analysis pipeline on a PDF.

    Returns a dict suitable for JSON serialization:
    {
        "summary": str,
        "keyPoints": [str, ...],
        "title": str,
        "wordCount": int,
        "pageCount": int,
        "readingTime": str,
        "sentiment": "positive" | "neutral" | "negative",
        "confidence": float (0-100),
    }
    """
    logger.info("=" * 50)
    logger.info("Starting AI analysis pipeline")

    # 1. Extract text
    logger.info("Extracting PDF text...")
    text, page_count, ocr_used = extract_pdf_text(pdf_bytes)

    if not text.strip():
        return {
            "summary": "No extractable text found in this PDF. The document may be scanned/image-based.",
            "keyPoints": ["No text content to analyze."],
            "title": "Untitled Document",
            "wordCount": 0,
            "pageCount": page_count,
            "readingTime": "< 1 minute",
            "sentiment": "neutral",
            "confidence": 0,
            "ocrUsed": ocr_used,
        }

    word_count = _count_words(text)
    logger.info("Word count: %d", word_count)

    # 2. Summarize
    logger.info("Generating summary...")
    summary = _generate_summary(text)

    # 3. Key points
    logger.info("Extracting key points...")
    key_points = _extract_key_points(text)

    # 4. Title
    logger.info("Generating title...")
    title = _generate_title(text)

    # 5. Sentiment
    logger.info("Analyzing sentiment...")
    sentiment_label, confidence = _analyze_sentiment(text)

    # 6. Reading time
    reading_time = _compute_reading_time(word_count)

    result = {
        "summary": summary,
        "keyPoints": key_points,
        "title": title,
        "wordCount": word_count,
        "pageCount": page_count,
        "readingTime": reading_time,
        "sentiment": sentiment_label,
        "confidence": confidence,
        "ocrUsed": ocr_used,
    }

    logger.info("AI analysis completed successfully")
    logger.info("=" * 50)
    return result


# ── Report generation (DOCX) ─────────────────────────────────────────

def generate_report(
    analysis_result: Dict,
    original_filename: str = "document.pdf",
) -> bytes:
    """Generate a formatted DOCX report from AI analysis results.

    Returns the DOCX file as bytes.
    """
    doc = DocxDocument()

    # -- Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Title (bold, no "AI Analysis Report" chrome — clean like iLovePDF)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(analysis_result.get("title", "Untitled"))
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    # -- Structured summary body. Each line is either "Label: content" (label
    #    bolded inline) or a "- item" list entry. Prefer the structured summary;
    #    fall back to the plain prose summary so we always render something.
    body_text = (
        analysis_result.get("structuredSummary")
        or analysis_result.get("summary")
        or "No summary available."
    )

    for raw_line in body_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        stripped = line.lstrip()
        # List item: "- foo" / "* foo" / "• foo"
        if stripped[:2] in ("- ", "* ") or stripped.startswith("•"):
            item = stripped.lstrip("-*• ").strip()
            doc.add_paragraph(item, style="List Bullet")
            continue

        # "Label: content" — bold the label portion.
        para = doc.add_paragraph()
        if ":" in stripped:
            label, _, content = stripped.partition(":")
            # Only treat as a label when it's short (avoid bolding sentences
            # that merely contain a colon).
            if len(label) <= 40:
                run = para.add_run(f"{label.strip()}: ")
                run.bold = True
                if content.strip():
                    para.add_run(content.strip())
                continue
        para.add_run(stripped)

    doc.add_paragraph()

    # -- Disclaimer (matches the honest "AI-generated" note iLovePDF uses)
    disclaimer = doc.add_paragraph()
    dis_run = disclaimer.add_run(
        "Disclaimer: AI-generated content may have inaccuracies. "
        "Please double-check important details."
    )
    dis_run.italic = True
    dis_run.font.size = Pt(9)
    dis_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── Report generation (PDF) ──────────────────────────────────────────

def generate_report_pdf(
    analysis_result: Dict,
    original_filename: str = "document.pdf",
    pdf_bytes: Optional[bytes] = None,
) -> bytes:
    """Generate a formatted PDF report from AI analysis results.

    Builds the same clean DOCX report as :func:`generate_report` (bold title,
    structured labeled summary, disclaimer), then converts it to PDF with
    LibreOffice (already used elsewhere for Office→PDF) rather than hand-rolling
    a second PDF layout engine.

    The structured summary is generated lazily here: when the caller hasn't
    supplied ``structuredSummary`` but passes the source ``pdf_bytes``, we
    extract the text and build it now — so the extra Groq call happens only on
    report download, not on every analyze.

    Returns the PDF file as bytes. Blocking (subprocess) — call via run_blocking.
    """
    import subprocess
    import uuid as _uuid
    from app.utils.concurrency import resolve_libreoffice_path

    libreoffice_path = resolve_libreoffice_path()
    if not libreoffice_path:
        raise RuntimeError("LibreOffice is not installed on the server")

    # Lazily build the structured summary from the source PDF if not provided.
    if not analysis_result.get("structuredSummary") and pdf_bytes:
        try:
            text, _, _ = extract_pdf_text(pdf_bytes)
            if text.strip():
                analysis_result = {
                    **analysis_result,
                    "structuredSummary": _generate_structured_summary(text),
                }
        except Exception as exc:  # noqa: BLE001 — fall back to prose summary
            logger.warning("Lazy structured summary failed: %s", exc)

    docx_bytes = generate_report(analysis_result, original_filename)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        input_path = temp_dir_path / "report.docx"
        input_path.write_bytes(docx_bytes)

        # Unique per-request profile dir so parallel LibreOffice runs don't
        # collide on the shared default user profile lock.
        profile_dir = temp_dir_path / f"lo_profile_{_uuid.uuid4().hex[:8]}"
        cmd = [
            libreoffice_path,
            "--headless",
            f'-env:UserInstallation=file:///{profile_dir.as_posix().lstrip("/")}',
            "--convert-to", "pdf",
            "--outdir", str(temp_dir_path),
            str(input_path),
        ]

        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if proc.returncode != 0:
            logger.error("AI report PDF conversion failed: %s", (proc.stderr or "")[:500])
            raise RuntimeError("PDF conversion failed")

        output_path = temp_dir_path / "report.pdf"
        if not output_path.exists():
            logger.error("AI report PDF conversion produced no output file")
            raise RuntimeError("PDF conversion produced no output")

        return output_path.read_bytes()